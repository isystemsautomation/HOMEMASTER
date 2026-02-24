"""
Generate product datasheet DOCX for HomeMaster modules.
Uses only information from README (or product page when available).
Include only specs/rows that have a value; omit anything not in source.
Output: <module>/Manuals/<MODULE_CODE>_Datasheet.docx
MiniPLC is never modified.
"""
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    subprocess.check_call(["pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

REPO_ROOT = Path(__file__).resolve().parent.parent

# Modules to generate (MiniPLC excluded)
MODULES = [
    "ALM-173-R1", "AIO-422-R1", "DIM-420-R1", "DIO-430-R1", "ENM-223-R1",
    "MicroPLC", "RGB-621-R1", "STR-3221-R1", "WLD-521-R1",
]

# Default publisher (MiniPLC block) â€” use unless README contradicts
PUBLISHER_DEFAULT = """ISYSTEMS AUTOMATION S.R.L.
Deligintei 18
Ploiesti
Romania
Tel: +40721389963
http://www.home-master.eu"""

# Internal: skip rows with no value from README (do not write "Not specified" in output)
NSR = None


def strip_md(s):
    if not s:
        return ""
    s = re.sub(r'!\[[^\]]*\]\([^)]+\)', '', s)
    s = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', s)
    s = re.sub(r'\*\*([^*]+)\*\*', r'\1', s)
    return s.strip()


def parse_table(lines, start_i):
    """Return (list of rows, next_index). Rows are list of cell strings."""
    rows = []
    i = start_i
    while i < len(lines):
        line = lines[i]
        s = line.strip()
        if not s or not s.startswith('|'):
            break
        if re.match(r'^\|[\s\-:|]+\|$', s):
            i += 1
            continue
        cells = [strip_md(c) for c in s.split('|')[1:-1]]
        if cells:
            rows.append(cells)
        i += 1
    return rows, i


def extract_module_code_and_name(readme_text, folder_name=None):
    """Get module code (e.g. ALM-173-R1) and product name from first # heading or README."""
    for line in readme_text.split('\n'):
        s = line.strip()
        if s.startswith('# ') and 'ðŸš§' not in s and 'Project Status' not in s:
            title = strip_md(s[2:].strip())
            m = re.search(r'([A-Z]{2,4}-[\d]+-R\d+|MicroPLC|MiniPLC)', title, re.I)
            code = m.group(1) if m else title.split()[0] if title else None
            if code:
                return code, title
            return (folder_name or "TBD"), title
    # Fallback: look for module code anywhere (e.g. Homemaster-MicroPLC)
    m = re.search(r'(MicroPLC|MiniPLC|[A-Z]{2,4}-\d+-R\d+)', readme_text, re.I)
    code = m.group(1) if m else (folder_name or "TBD")
    return code, (code if isinstance(code, str) else "Product")


def extract_product_description(readme_text):
    """1â€“2 paragraphs: what it is, what it connects to, intended use. From Overview/Introduction."""
    lines = readme_text.replace('\r\n', '\n').split('\n')
    paras = []
    in_intro = False
    for i, line in enumerate(lines):
        s = line.strip()
        if re.match(r'^#+\s*(1\.\s*)?(Overview|Introduction|Module Description|Overview & Description)', s, re.I):
            in_intro = True
            continue
        if in_intro:
            if s.startswith('#') and not s.startswith('## '):
                break
            if s.startswith('## ') and 'Core' not in s and 'System Role' not in s:
                break
            if s.startswith('**') and s.endswith('**') and len(s) < 80:
                continue
            if s.startswith('>'):
                s = s.lstrip('>').strip()
            if s and not s.startswith('|') and not s.startswith('!['):
                p = strip_md(s)
                if len(p) > 30 and p not in [strip_md(x) for x in paras]:
                    paras.append(p)
                    if len(paras) >= 2:
                        break
    if not paras:
        # Fallback: first substantial paragraph after first # heading
        for line in lines:
            p = strip_md(line.strip())
            if len(p) > 50 and not p.startswith('http'):
                paras.append(p)
                break
    return ' '.join(paras[:2]) if paras else ""


def extract_technical_specs(readme_text):
    """Build list of (Specification, Details) from README tables and bullets. No guessing."""
    lines = readme_text.replace('\r\n', '\n').split('\n')
    specs = []
    seen = set()

    # Find 2-col or 3-col tables that look like spec tables (Parameter|Min|Typ|Max or Spec|Details)
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith('|'):
            rows, j = parse_table(lines, i)
            i = j
            if len(rows) < 2:
                continue
            headers = [h.lower() for h in rows[0]]
            # Key ratings / Electrical: Parameter, Min, Typ, Max, Notes
            if 'parameter' in str(headers) or 'specification' in str(headers):
                for r in rows[1:]:
                    if len(r) >= 2:
                        spec_name = r[0].strip()
                        if not spec_name or spec_name in seen:
                            continue
                        seen.add(spec_name)
                        # Use "Details" column or concat Min/Typ/Max/Notes
                        if len(r) >= 5 and any(x for x in r[1:5]):
                            details = ' / '.join(x for x in r[1:5] if x.strip())
                        else:
                            details = (r[1] if len(r) > 1 else "").strip()
                        if not details:
                            continue
                        specs.append((spec_name, details))
            # Subsystem / Interface table: Subsystem | Qty | Description
            if 'subsystem' in str(headers) or 'interface' in str(headers):
                for r in rows[1:]:
                    if len(r) >= 2:
                        name = r[0].strip()
                        if not name or name in seen:
                            continue
                        seen.add(name)
                        val = r[1].strip() if len(r) > 1 else ''
                        desc = r[2].strip() if len(r) > 2 else ''
                        details = f"{val}; {desc}" if desc else val
                        if not details:
                            continue
                        specs.append((name, details))
        i += 1

    # Explicit power / supply from text (only if stated)
    if not any('supply' in s[0].lower() or 'power' in s[0].lower() for s in specs):
        if re.search(r'24\s*VDC|24\s*V\s*DC|24\s*V\s*DC', readme_text, re.I):
            specs.insert(0, ("Power supply", "24 V DC (SELV)"))

    # MCU if mentioned
    mcu = re.search(r'(RP2350|ESP32|RP2350A)[^\n]*', readme_text)
    if mcu and not any('MCU' in s[0] or 'microcontroller' in s[0].lower() for s in specs):
        specs.append(("MCU", strip_md(mcu.group(0))[:80]))

    return specs


def extract_installation_mechanical(readme_text):
    """3-column table: Category | Specification | Details. Only rows where README states a value."""
    rows = []
    term = re.search(r'terminal[^\n]*(5\.08|pitch|wire|torque|Nm|mmÂ²|AWG)[^\n]*', readme_text, re.I | re.DOTALL)
    if term:
        rows.append(("Terminal Specifications", "Terminal type / pitch", strip_md(term.group(0))[:120]))
    op_temp = re.search(r'(operating\s*temp|0\s*Â°?C\s*to\s*\+?40\s*Â°?C|0\.\.\.?\+?40)', readme_text, re.I)
    if op_temp:
        rows.append(("Environmental Ratings", "Operating temperature", op_temp.group(0).strip()))
    ip_ = re.search(r'IP20|IP\s*20|ingress\s*protection', readme_text, re.I)
    if ip_:
        rows.append(("Environmental Ratings", "Ingress protection", ip_.group(0)))
    humidity = re.search(r'(\d+\s*%?\s*RH|humidity[^\n]{0,30})', readme_text, re.I)
    if humidity:
        rows.append(("Environmental Ratings", "Relative humidity", strip_md(humidity.group(0))))
    dims = re.search(r'(\d+\.?\d*\s*Ã—\s*\d+\.?\d*\s*Ã—\s*\d+\.?\d*\s*mm|\d+\.?\d*\s*Ã—\s*\d+\.?\d*\s*mm)', readme_text)
    if dims:
        rows.append(("Mechanical & Packaging", "Product dimensions", dims.group(0)))
    din = re.search(r'(35\s*mm\s*DIN|DIN\s*rail|EN\s*50022)', readme_text, re.I)
    if din:
        rows.append(("Mechanical & Packaging", "Mounting", din.group(0)))
    weight = re.search(r'(\d+\s*g|\d+\s*kg|weight[^\n]{0,20})', readme_text, re.I)
    if weight:
        rows.append(("Mechanical & Packaging", "Net weight", weight.group(0)))
    return rows


def extract_compliance_and_safety(readme_text):
    """Compliance: only if README states. Safety notice only if SELV/mains/isolation mentioned."""
    compliance = []
    if re.search(r'CE\b|RoHS|REACH|UL\d|CSA|directive|certification|compliant', readme_text, re.I):
        for m in re.finditer(r'([CE]E|RoHS|REACH|UL[\d\-]+|CSA[^\n]{0,30}|compliant[^\n]{0,40})', readme_text, re.I):
            compliance.append(strip_md(m.group(0)))
    safety = None
    if re.search(r'SELV|PELV|mains|hazardous\s*voltage|qualified\s*personnel|isolation\s*domain', readme_text, re.I):
        safety = "Safety notice: This module is intended for SELV/PELV applications. Install and wire according to local regulations. Do not connect mains to logic/input terminals unless so specified."
    return compliance, safety


def extract_publisher(readme_text):
    """If README has manufacturer/publisher, use it; else default."""
    if re.search(r'home-master\.eu|home-master\.eu|ISYSTEMS|Ploiesti|Romania', readme_text, re.I):
        return PUBLISHER_DEFAULT
    return PUBLISHER_DEFAULT


def build_docx(module_code, product_name, part_no, readme_text, out_path):
    doc = Document()

    # A) Top section
    p = doc.add_paragraph()
    p.add_run("Datasheet").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(product_name or module_code)
    doc.add_paragraph(f"Part No. {part_no or 'TBD'}")

    # B) Product Description (only from README)
    doc.add_heading("Product Description", level=1)
    desc = extract_product_description(readme_text)
    if desc:
        doc.add_paragraph(desc)
    else:
        doc.add_paragraph(f"{product_name or module_code}. See module README for full description.")

    # C) Technical Specifications â€” only rows with value from README
    doc.add_heading("Technical Specifications", level=1)
    specs = extract_technical_specs(readme_text)
    if specs:
        table = doc.add_table(rows=1 + len(specs), cols=2)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = "Specification"
        table.rows[0].cells[1].text = "Details"
        for i, (spec, detail) in enumerate(specs):
            if detail:
                table.rows[i + 1].cells[0].text = spec[:200]
                table.rows[i + 1].cells[1].text = detail[:400]
    else:
        doc.add_paragraph("See module README for technical specifications.")

    # D) Installation, Environmental & Mechanical â€” only rows with value from README
    doc.add_heading("Installation, Environmental & Mechanical", level=1)
    inst_rows = extract_installation_mechanical(readme_text)
    if inst_rows:
        t2 = doc.add_table(rows=1 + len(inst_rows), cols=3)
        t2.style = 'Table Grid'
        t2.rows[0].cells[0].text = "Category"
        t2.rows[0].cells[1].text = "Specification"
        t2.rows[0].cells[2].text = "Details"
        for i, row in enumerate(inst_rows):
            t2.rows[i + 1].cells[0].text = row[0][:100]
            t2.rows[i + 1].cells[1].text = row[1][:100]
            t2.rows[i + 1].cells[2].text = (row[2] or "")[:200]
    else:
        doc.add_paragraph("See module README for installation and mechanical data.")

    # E) Compliance & Certifications â€” only if README states something
    doc.add_heading("Compliance & Certifications", level=1)
    compliance, safety = extract_compliance_and_safety(readme_text)
    if compliance:
        for c in compliance:
            doc.add_paragraph(c)
    if safety:
        doc.add_paragraph(safety)
    if not compliance and not safety:
        doc.add_paragraph("See module README for compliance information.")

    # F) Published by
    doc.add_heading("Published by", level=1)
    doc.add_paragraph(extract_publisher(readme_text))

    doc.save(out_path)


def process_module(module_name, part_no=None):
    if module_name == "MiniPLC":
        print("  Skip MiniPLC: datasheet is ready, not modified.")
        return False
    module_path = REPO_ROOT / module_name
    readme_path = module_path / "README.md"
    manuals_dir = module_path / "Manuals"
    if not readme_path.exists():
        print(f"  Skip {module_name}: no README.md")
        return False
    readme_text = readme_path.read_text(encoding='utf-8', errors='replace')
    module_code, product_name = extract_module_code_and_name(readme_text, folder_name=module_name)
    out_name = f"{module_code}_Datasheet.docx"
    out_path = manuals_dir / out_name
    manuals_dir.mkdir(parents=True, exist_ok=True)
    build_docx(module_code, product_name, part_no or "TBD", readme_text, out_path)
    print(f"  OK {module_name} -> {out_path.relative_to(REPO_ROOT)}")
    return True


def main():
    print("Generating module datasheets (MiniPLC structure, README-only content)...")
    for m in MODULES:
        process_module(m)
    print("Done.")


if __name__ == "__main__":
    main()
