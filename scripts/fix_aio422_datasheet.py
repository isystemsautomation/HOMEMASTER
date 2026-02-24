"""
Fix AIO-422-R1_Datasheet.docx: README primary, PDF secondary.
Only include info that exists in README or PDF; omit rows/sections with no source.
"""
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    subprocess.check_call(["pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

REPO_ROOT = Path(__file__).resolve().parent.parent
MANUALS = REPO_ROOT / "AIO-422-R1" / "Manuals"
OUT_PATH = MANUALS / "AIO-422-R1_Datasheet.docx"

# From PDF (previous datasheet) - user-provided
PDF_PART_NO = "0000007"
PDF_PUBLISHER = """ISYSTEMS AUTOMATION S.R.L.
Deligintei 18
Ploiesti
Romania
Tel: +40721389963
http://www.home-master.eu"""


def build_fixed_docx():
    doc = Document()

    # A) Title block
    p = doc.add_paragraph()
    p.add_run("Datasheet").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("AIO-422-R1 Module")
    doc.add_paragraph(f"Part No. {PDF_PART_NO}")

    # B) Product Description (from README)
    doc.add_heading("Product Description", level=1)
    doc.add_paragraph(
        "The AIO-422-R1 is a high-precision analog I/O expansion module designed for home automation, "
        "HVAC, and environmental monitoring. It connects to MicroPLC or MiniPLC controllers via RS-485 (Modbus RTU) "
        "and integrates with ESPHome and Home Assistant for analog sensing and control in smart automation systems."
    )
    doc.add_paragraph(
        "Typical use includes real-time analog signal monitoring, analog output automation (lighting, HVAC, pumps), "
        "accurate temperature-based triggers using RTDs, and custom automation logic via MQTT or native ESPHome entities."
    )

    # C) Technical Specifications (README + PDF only; no placeholders)
    doc.add_heading("Technical Specifications", level=1)
    spec_rows = [
        ("Power supply", "24 V DC"),
        ("MCU", "RP2350A (from previous datasheet)."),
        ("Analog inputs", "4 × 0–10 V (ADS1115, 16-bit)"),
        ("Analog outputs", "2 × 0–10 V (MCP4725, 12-bit)"),
        ("RTD inputs", "2 × PT100/PT1000 (MAX31865); 2-, 3-, or 4-wire"),
        ("Input/Output protection", "ESD, overvoltage"),
        ("Communication interface", "RS-485 (Modbus RTU); half-duplex, up to 115.2 kbps, with protections (from previous datasheet)."),
        ("USB port", "USB Type-C; USB 2.0, CP2102N, ESD protection (from previous datasheet)."),
        ("RTC", "PCF8563 (from previous datasheet)."),
        ("User interface", "4 front-panel buttons, status LEDs"),
        ("Dimensions", "DIN-rail, 3 modules wide; 52.5 × 90 × 59 mm, 9 division units (from previous datasheet)."),
        ("Compatibility", "MicroPLC, MiniPLC"),
    ]
    table = doc.add_table(rows=1 + len(spec_rows), cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Specification"
    table.rows[0].cells[1].text = "Details"
    for i, (spec, detail) in enumerate(spec_rows):
        table.rows[i + 1].cells[0].text = spec
        table.rows[i + 1].cells[1].text = detail

    # D) Installation, Environmental & Mechanical (PDF structure; README has no details)
    doc.add_heading("Installation, Environmental & Mechanical", level=1)
    inst_rows = [
        ("Terminal Specifications", "Terminal type / pitch", "5.08 mm pitch (from previous datasheet)."),
        ("Terminal Specifications", "Wire cross-section / rating", "300 V, 20 A; 26–12 AWG (from previous datasheet)."),
        ("Terminal Specifications", "Tightening torque", "0.5–0.6 Nm (from previous datasheet)."),
        ("Environmental Ratings", "Operating temperature", "0 °C to +40 °C (from previous datasheet)."),
        ("Environmental Ratings", "Relative humidity", "Max 95 % RH, non-condensing (from previous datasheet)."),
        ("Environmental Ratings", "Ingress protection", "IP20 (from previous datasheet)."),
        ("Mechanical & Packaging", "Product dimensions", "52.5 × 90 × 59 mm (from previous datasheet)."),
        ("Mechanical & Packaging", "DIN units", "9 division units (from previous datasheet)."),
        ("Mechanical & Packaging", "Mounting", "DIN rail per EN 50022, 35 mm (from previous datasheet)."),
    ]
    t2 = doc.add_table(rows=1 + len(inst_rows), cols=3)
    t2.style = "Table Grid"
    t2.rows[0].cells[0].text = "Category"
    t2.rows[0].cells[1].text = "Specification"
    t2.rows[0].cells[2].text = "Details"
    for i, row in enumerate(inst_rows):
        t2.rows[i + 1].cells[0].text = row[0]
        t2.rows[i + 1].cells[1].text = row[1]
        t2.rows[i + 1].cells[2].text = row[2]

    # E) Compliance & Certifications (from PDF; replace corrupted "ce" content)
    doc.add_heading("Compliance & Certifications", level=1)
    doc.add_paragraph("IP20.")
    doc.add_paragraph("Operation Type 1 (UL60730-1, CSA E60730-1).")
    doc.add_paragraph("Rated impulse voltage for digital output: 2.5 kV (UL60730-1, altitude ≤ 2000 m).")
    doc.add_paragraph("Degree of pollution: 2.")
    doc.add_paragraph("Certifications: CE, UL60730-1, CSA E60730-1 (from previous datasheet).")

    # F) Published by
    doc.add_heading("Published by", level=1)
    doc.add_paragraph(PDF_PUBLISHER)

    MANUALS.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    build_fixed_docx()
