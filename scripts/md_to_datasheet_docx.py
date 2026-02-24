"""
Generate Datasheet.docx for each HomeMaster module (except MiniPLC) using the SAME
structure and sections as the MiniPLC datasheet. Content is taken from each module's README.
MiniPLC datasheet is left unchanged.
Output: <module>/Manuals/Datasheet.docx
Usage: python md_to_datasheet_docx.py
"""
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    import subprocess
    subprocess.check_call(["pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches

REPO_ROOT = Path(__file__).resolve().parent.parent

# Modules to generate datasheets for — MiniPLC is EXCLUDED (its datasheet is ready, do not touch)
MODULES = [
    "ALM-173-R1", "AIO-422-R1", "DIM-420-R1", "DIO-430-R1", "ENM-223-R1",
    "MicroPLC", "RGB-621-R1", "STR-3221-R1", "WLD-521-R1",
]

# Exact section order as MiniPLC datasheet — nothing in plus
MINIPLC_SECTIONS = [
    "Overview",
    "Quick Overview",
    "Typical Applications",
    "Tech Specs",
    "Installation, Environmental & Mechanical",
    "Home Assistant & ESPHome Integration",
    "Documentation",
    "Power Supply",
    "Inputs & Outputs",
    "Cable Recommendations & Shield Grounding",
    "System Architecture & Pinout",
    "ESPHome YAML Configuration & Hardware Pinout Manual",
    "Links",
]

# Map README section titles (normalized) to our section name
SECTION_ALIASES = {
    "overview": ["overview", "introduction", "module description", "1. introduction", "1 introduction", "introduction", "1. overview & description", "overview & description"],
    "quick overview": ["quick overview", "highlights", "2.2 quick overview", "core capabilities", "overview of the", "2.2 overview"],
    "typical applications": ["typical applications", "use cases", "2. use cases", "use case", "example use cases"],
    "tech specs": [
        "tech specs", "technical specification", "i/o summary", "electrical specifications", "2.2 i/o",
        "2.3 i/o", "2.5", "2.6 key ratings", "absolute ratings", "technical specifications",
        "2.2 overview", "2.3 i/o summary", "2.4 terminals", "2.5 electrical", "2.6 reliability",
        "specifications", "key features", "⚙️ technical specifications",
    ],
    "installation, environmental & mechanical": [
        "installation, environmental & mechanical", "mechanical", "2.6 mechanical", "2.7 mechanical",
        "mechanical & environmental", "2.4 terminals", "environmental & compliance", "2.7 mechanical & compliance",
        "2.6 mechanical & environmental", "mechanical & compliance", "5.6 protection", "5.7 mechanical",
        "5.8 environmental",
    ],
    "home assistant & esphome integration": [
        "home assistant & esphome integration", "esphome integration", "7. esphome integration guide",
        "installation & quick start", "4. installation & quick start", "5. installation & quick start",
        "getting started", "quick start",
    ],
    "documentation": [
        "documentation", "downloads", "11. downloads", "12. downloads", "hardware design files",
        "firmware & software", "manuals & datasheets",
    ],
    "power supply": [
        "power supply", "power", "5.2 power", "4.2 power", "4.2.1 supply", "supply", "power input",
        "4.3 networking", "supply details", "supply types",
    ],
    "inputs & outputs": [
        "inputs & outputs", "i/o summary", "2.3 i/o", "2.4 terminals", "terminals & pinout",
        "digital inputs", "relay", "relays", "2.3 connectors", "2.4 reliability", "2.5 functional",
        "i/o and electrical summary", "terminals & pinout (field side)", "2.4 terminals & pinout",
    ],
    "cable recommendations & shield grounding": [
        "cable recommendations & shield grounding", "cable recommendations", "shielding", "shield grounding",
        "4.3 i/o & interface warnings", "shielding & emc",
    ],
    "system architecture & pinout": [
        "system architecture & pinout", "system architecture", "pinout", "2.1 diagrams",
        "diagrams & pinouts", "2.1 diagrams & pinouts", "5.1 diagrams & pinouts",
    ],
    "esphome yaml configuration & hardware pinout manual": [
        "esphome yaml configuration & hardware pinout manual", "esphome yaml", "programming",
        "8. programming & customization", "8. programming", "programming & customization",
        "pinout", "gpio pinout", "esp32 gpio pinout",
    ],
    "links": ["links", "support", "12. support", "11. support", "downloads"],
}


def strip_md(s):
    """Remove markdown images and trim."""
    s = re.sub(r'!\[[^\]]*\]\([^)]+\)', '', s)
    s = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', s)  # links -> text
    return s.strip()


def normalize_heading(title):
    """Normalize for matching: lowercase, remove extra spaces, strip leading numbers."""
    t = title.lower().strip()
    t = re.sub(r'^\d+[\.\)]\s*', '', t)  # "1. Introduction" -> "introduction"
    t = re.sub(r'^#+\s*', '', t)
    t = re.sub(r'\s+', ' ', t)
    return t.strip()


def is_table_sep(line):
    return bool(re.match(r'^\|[\s\-:|]+\|$', line.strip()))


def parse_table(lines, i):
    rows = []
    while i < len(lines):
        line = lines[i]
        s = line.strip()
        if not s:
            break
        if s.startswith('|') and is_table_sep(s):
            i += 1
            continue
        if s.startswith('|'):
            cells = [strip_md(c) for c in s.split('|')[1:-1]]
            if cells:
                rows.append(cells)
        else:
            break
        i += 1
    return rows, i


def parse_readme_into_sections(md_text):
    """
    Split README by ## and ###. Return list of (level, title, content_blocks).
    content_blocks: list of ('p', text) | ('table', rows) | ('list', text) | ('blockquote', text)
    """
    md_text = re.sub(r'^# 🚧.*?---\s*', '', md_text, flags=re.DOTALL)
    md_text = re.sub(r'^> \*\*Important Notice:\*\*.*?---\s*', '', md_text, flags=re.DOTALL)
    lines = md_text.replace('\r\n', '\n').split('\n')
    sections = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith('# '):
            title = strip_md(line[2:].strip())
            i += 1
            content = []
            while i < len(lines):
                cur = lines[i]
                if cur.startswith('# ') or (cur.startswith('## ') and not sections):
                    break
                if cur.startswith('## '):
                    break
                s = cur.strip()
                if not s:
                    i += 1
                    continue
                if s.startswith('### '):
                    content.append(('h3', strip_md(s[4:])))
                    i += 1
                    continue
                if s.startswith('#### '):
                    content.append(('h3', strip_md(s[5:])))
                    i += 1
                    continue
                if s.startswith('|'):
                    table_rows, i = parse_table(lines, i)
                    if table_rows:
                        content.append(('table', table_rows))
                    continue
                if s.startswith('>'):
                    content.append(('blockquote', strip_md(s.lstrip('>').strip())))
                    i += 1
                    continue
                if s.startswith('- ') or s.startswith('* '):
                    content.append(('list', strip_md(s[2:])))
                    i += 1
                    continue
                if re.match(r'^\d+\.\s', s):
                    content.append(('list', strip_md(re.sub(r'^\d+\.\s*', '', s))))
                    i += 1
                    continue
                content.append(('p', strip_md(s)))
                i += 1
            sections.append((1, title, content))
            continue
        if line.startswith('## '):
            title = strip_md(line[3:].strip())
            i += 1
            content = []
            while i < len(lines):
                cur = lines[i]
                if cur.startswith('# '):
                    break
                if cur.startswith('## '):
                    break
                s = cur.strip()
                if not s:
                    i += 1
                    continue
                if s.startswith('### '):
                    content.append(('h3', strip_md(s[4:])))
                    i += 1
                    continue
                if s.startswith('#### '):
                    content.append(('h3', strip_md(s[5:])))
                    i += 1
                    continue
                if s.startswith('|'):
                    table_rows, i = parse_table(lines, i)
                    if table_rows:
                        content.append(('table', table_rows))
                    continue
                if s.startswith('>'):
                    content.append(('blockquote', strip_md(s.lstrip('>').strip())))
                    i += 1
                    continue
                if s.startswith('- ') or s.startswith('* '):
                    content.append(('list', strip_md(s[2:])))
                    i += 1
                    continue
                if re.match(r'^\d+\.\s', s):
                    content.append(('list', strip_md(re.sub(r'^\d+\.\s*', '', s))))
                    i += 1
                    continue
                content.append(('p', strip_md(s)))
                i += 1
            sections.append((2, title, content))
            continue
        i += 1
    return sections


def find_content_for_section(sections, target_section):
    """Find content from sections that best matches target_section (one of MINIPLC_SECTIONS)."""
    target_norm = target_section.lower()
    aliases = SECTION_ALIASES.get(target_norm, [target_norm])
    for level, title, content in sections:
        norm = normalize_heading(title)
        for a in aliases:
            if a in norm or norm in a:
                return content
        if target_norm in norm:
            return content
    return None


def add_blocks_to_doc(doc, blocks):
    """Append content blocks to document."""
    if not blocks:
        return
    for item in blocks:
        if item[0] == 'h3':
            doc.add_heading(item[1], level=2)
        elif item[0] == 'p':
            if item[1]:
                doc.add_paragraph(item[1])
        elif item[0] == 'blockquote':
            if item[1]:
                p = doc.add_paragraph(item[1])
                p.paragraph_format.left_indent = Inches(0.25)
        elif item[0] == 'list':
            if item[1]:
                doc.add_paragraph(item[1], style='List Bullet')
        elif item[0] == 'table':
            rows = item[1]
            if not rows:
                continue
            cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = 'Table Grid'
            for r, row in enumerate(rows):
                for c, cell in enumerate(row):
                    if c < cols:
                        table.rows[r].cells[c].text = cell[:600]
            doc.add_paragraph('')


def build_datasheet_docx(module_name):
    """Build datasheet with MiniPLC structure, content from module README. MiniPLC is never modified."""
    if module_name == "MiniPLC":
        print("  Skip MiniPLC: datasheet is ready, not modified.")
        return False
    module_path = REPO_ROOT / module_name
    readme_path = module_path / "README.md"
    manuals_dir = module_path / "Manuals"
    out_path = manuals_dir / "Datasheet.docx"

    if not readme_path.exists():
        print(f"  Skip {module_name}: no README.md")
        return False

    try:
        md_text = readme_path.read_text(encoding='utf-8', errors='replace')
    except Exception as e:
        print(f"  Error reading {module_name}/README.md: {e}")
        return False

    sections = parse_readme_into_sections(md_text)
    if not sections:
        print(f"  Skip {module_name}: no sections parsed")
        return False

    # Title from first H1
    doc = Document()
    first_h1 = None
    for lev, title, _ in sections:
        if lev == 1 and title:
            first_h1 = title
            break
    if not first_h1:
        for lev, title, _ in sections:
            if title and not title.startswith('#'):
                first_h1 = title
                break
    if first_h1:
        doc.add_heading(first_h1, level=0)
        doc.core_properties.title = f"{first_h1} — Datasheet"
    else:
        doc.add_heading(module_name, level=0)

    # Tagline (from first paragraph or "DIN rail • ..." style)
    tagline = "DIN rail • 24 V DC • Modbus RTU • Home Assistant"
    doc.add_paragraph(tagline)

    # Exactly these sections in order
    for sec in MINIPLC_SECTIONS:
        content = find_content_for_section(sections, sec)
        doc.add_heading(sec, level=1)
        if content:
            add_blocks_to_doc(doc, content)
        else:
            doc.add_paragraph("See module README for details.")

    manuals_dir.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"  OK {module_name} -> {out_path.relative_to(REPO_ROOT)}")
    return True


def main():
    print("Generating Datasheet.docx (MiniPLC structure) from README for each module...")
    for module in MODULES:
        build_datasheet_docx(module)
    print("Done.")


if __name__ == "__main__":
    main()
